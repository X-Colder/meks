import uuid
from html import escape
from io import BytesIO
from urllib.parse import quote

from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from meks.api.schemas.paper import (
    BlockCreate,
    BlockResponse,
    BlockUpdate,
    ImportBlockRequest,
    PaperExportRequest,
    PaperCreate,
    PaperDetailResponse,
    PaperResponse,
    PaperUpdate,
    ReorderBlocksRequest,
)
from meks.core.exceptions import ForbiddenException, NotFoundException
from meks.core.rbac import Permission
from meks.dependencies import require_permission
from meks.models.base import get_db
from meks.models.paper import BlockType, Paper, PaperBlock, PaperStatus
from meks.models.user import User

router = APIRouter()


def _download_filename(name: str, ext: str) -> str:
    safe_name = "".join(ch if ch.isalnum() or ch in "-_." else "_" for ch in name).strip("_")
    safe_name = safe_name[:80] or "paper"
    return f"{safe_name}.{ext}"


def _plain_markdown_line(line: str) -> str:
    line = line.strip()
    while line.startswith("#"):
        line = line[1:].strip()
    for token in ("**", "__", "`"):
        line = line.replace(token, "")
    return line


def _paper_markdown(paper: Paper, blocks: list[PaperBlock]) -> str:
    doc_block = next((b for b in blocks if b.block_type == BlockType.text and not b.source_type), None)
    if doc_block and doc_block.content.strip():
        return doc_block.content.strip()
    parts = [f"# {paper.title}"]
    if paper.abstract:
        parts.extend(["", "## 摘要", paper.abstract])
    return "\n".join(parts)


def _add_docx_watermark(section, watermark_text: str) -> None:
    from docx.oxml import parse_xml

    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    escaped = escape(watermark_text)
    positions = [
        (-120, 120),
        (120, 120),
        (-120, 330),
        (120, 330),
        (-120, 540),
        (120, 540),
    ]
    for idx, (left, top) in enumerate(positions):
        watermark = parse_xml(f"""
        <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:v="urn:schemas-microsoft-com:vml"
             xmlns:o="urn:schemas-microsoft-com:office:office">
          <w:pict>
            <v:shape id="MEKSWatermark{idx}" o:spid="_x0000_s{2048 + idx}" type="#_x0000_t136"
              style="position:absolute;margin-left:{left}pt;margin-top:{top}pt;width:360pt;height:80pt;rotation:315;z-index:-251654144;mso-position-horizontal-relative:page;mso-position-vertical-relative:page"
              fillcolor="#d9d9d9" stroked="f">
              <v:fill opacity="0.18"/>
              <v:textpath style="font-family:Arial;font-size:28pt" string="{escaped}"/>
            </v:shape>
          </w:pict>
        </w:r>
        """)
        paragraph._p.append(watermark)


def _build_word_export(paper: Paper, blocks: list[PaperBlock], watermark_text: str | None) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    if watermark_text:
        _add_docx_watermark(doc.sections[0], watermark_text)

    title = doc.add_heading(paper.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    content = _paper_markdown(paper, blocks)
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(_plain_markdown_line(line), level=1)
        elif line.startswith("## "):
            doc.add_heading(_plain_markdown_line(line), level=2)
        elif line.startswith("### "):
            doc.add_heading(_plain_markdown_line(line), level=3)
        elif line[:3].lstrip().startswith(("-", "*")):
            doc.add_paragraph(_plain_markdown_line(line[1:]), style="List Bullet")
        elif len(line) > 2 and line[0].isdigit() and line[1] == ".":
            doc.add_paragraph(_plain_markdown_line(line[2:]), style="List Number")
        else:
            paragraph = doc.add_paragraph(_plain_markdown_line(line))
            paragraph.paragraph_format.first_line_indent = Pt(21)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _build_pdf_export(paper: Paper, blocks: list[PaperBlock], watermark_text: str | None) -> bytes:
    import textwrap
    import fitz

    pdf = fitz.open()
    width, height = fitz.paper_size("a4")
    margin = 56
    line_height = 18
    y = margin
    page = pdf.new_page(width=width, height=height)

    def add_watermark(target_page):
        if not watermark_text:
            return
        for x in (95, width / 2 - 110, width - 315):
            for y_pos in (150, height / 2, height - 170):
                point = fitz.Point(x, y_pos)
                target_page.insert_text(
                    point,
                    watermark_text,
                    fontsize=32,
                    fontname="china-s",
                    color=(0.78, 0.78, 0.78),
                    fill_opacity=0.18,
                    morph=(point, fitz.Matrix(1, 1).prerotate(45)),
                )

    def ensure_space(lines: int = 1):
        nonlocal page, y
        if y + lines * line_height > height - margin:
            add_watermark(page)
            page = pdf.new_page(width=width, height=height)
            y = margin

    def write_line(text: str, size: int = 11, bold: bool = False, indent: int = 0):
        nonlocal y
        ensure_space()
        fontname = "china-s"
        page.insert_text(
            (margin + indent, y),
            text,
            fontsize=size,
            fontname=fontname,
            color=(0.08, 0.08, 0.08),
        )
        y += line_height + (4 if bold else 0)

    write_line(paper.title, size=18, bold=True)
    y += 8

    content = _paper_markdown(paper, blocks)
    for raw_line in content.splitlines():
        line = _plain_markdown_line(raw_line)
        if not line:
            y += 6
            continue
        size = 14 if raw_line.strip().startswith("#") else 11
        indent = 0 if size > 11 else 18
        max_chars = 34 if size > 11 else 46
        for wrapped in textwrap.wrap(line, width=max_chars, break_long_words=False, replace_whitespace=False):
            write_line(wrapped, size=size, bold=size > 11, indent=indent)

    add_watermark(page)
    output = pdf.tobytes()
    pdf.close()
    return output


async def _get_paper_for_user(paper_id: str, user: User, db: AsyncSession) -> Paper:
    result = await db.execute(select(Paper).where(Paper.id == uuid.UUID(paper_id)))
    paper = result.scalar_one_or_none()
    if not paper:
        raise NotFoundException("论文")
    if paper.owner_id != user.id:
        raise ForbiddenException("只有论文创建者可以访问此论文")
    return paper


@router.post("", response_model=PaperResponse)
async def create_paper(
    request: PaperCreate,
    user: User = Depends(require_permission(Permission.DOC_READ)),
    db: AsyncSession = Depends(get_db),
):
    paper = Paper(
        title=request.title,
        abstract=request.abstract,
        target_journal=request.target_journal,
        owner_id=user.id,
    )
    db.add(paper)
    await db.commit()
    await db.refresh(paper)
    return paper


@router.get("", response_model=list[PaperResponse])
async def list_papers(
    user: User = Depends(require_permission(Permission.DOC_READ)),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Paper).where(Paper.owner_id == user.id).order_by(Paper.created_at.desc())
    )
    return result.scalars().all()


@router.get("/{paper_id}", response_model=PaperDetailResponse)
async def get_paper(
    paper_id: str,
    user: User = Depends(require_permission(Permission.DOC_READ)),
    db: AsyncSession = Depends(get_db),
):
    paper = await _get_paper_for_user(paper_id, user, db)
    blocks_result = await db.execute(
        select(PaperBlock)
        .where(PaperBlock.paper_id == paper.id)
        .order_by(PaperBlock.sort_order)
    )
    blocks = blocks_result.scalars().all()
    response = PaperDetailResponse.model_validate(paper)
    response.blocks = [BlockResponse.model_validate(b) for b in blocks]
    return response


@router.patch("/{paper_id}", response_model=PaperResponse)
async def update_paper(
    paper_id: str,
    request: PaperUpdate,
    user: User = Depends(require_permission(Permission.DOC_READ)),
    db: AsyncSession = Depends(get_db),
):
    paper = await _get_paper_for_user(paper_id, user, db)
    update_data = request.model_dump(exclude_unset=True)
    if "status" in update_data:
        update_data["status"] = PaperStatus(update_data["status"])
    for field, value in update_data.items():
        setattr(paper, field, value)
    await db.commit()
    await db.refresh(paper)
    return paper


@router.delete("/{paper_id}")
async def delete_paper(
    paper_id: str,
    user: User = Depends(require_permission(Permission.DOC_READ)),
    db: AsyncSession = Depends(get_db),
):
    paper = await _get_paper_for_user(paper_id, user, db)
    await db.delete(paper)
    await db.commit()
    return {"detail": "论文已删除"}


@router.post("/{paper_id}/blocks", response_model=BlockResponse)
async def add_block(
    paper_id: str,
    request: BlockCreate,
    user: User = Depends(require_permission(Permission.DOC_READ)),
    db: AsyncSession = Depends(get_db),
):
    paper = await _get_paper_for_user(paper_id, user, db)
    block = PaperBlock(
        paper_id=paper.id,
        block_type=BlockType(request.block_type),
        content=request.content,
        sort_order=request.sort_order,
        source_type=request.source_type,
        source_id=request.source_id,
        extra=request.extra,
    )
    db.add(block)
    await db.commit()
    await db.refresh(block)
    return block


@router.patch("/{paper_id}/blocks/{block_id}", response_model=BlockResponse)
async def update_block(
    paper_id: str,
    block_id: str,
    request: BlockUpdate,
    user: User = Depends(require_permission(Permission.DOC_READ)),
    db: AsyncSession = Depends(get_db),
):
    await _get_paper_for_user(paper_id, user, db)
    result = await db.execute(
        select(PaperBlock).where(
            PaperBlock.id == uuid.UUID(block_id),
            PaperBlock.paper_id == uuid.UUID(paper_id),
        )
    )
    block = result.scalar_one_or_none()
    if not block:
        raise NotFoundException("内容块")
    update_data = request.model_dump(exclude_unset=True)
    if "block_type" in update_data:
        update_data["block_type"] = BlockType(update_data["block_type"])
    for field, value in update_data.items():
        setattr(block, field, value)
    await db.commit()
    await db.refresh(block)
    return block


@router.delete("/{paper_id}/blocks/{block_id}")
async def delete_block(
    paper_id: str,
    block_id: str,
    user: User = Depends(require_permission(Permission.DOC_READ)),
    db: AsyncSession = Depends(get_db),
):
    await _get_paper_for_user(paper_id, user, db)
    result = await db.execute(
        select(PaperBlock).where(
            PaperBlock.id == uuid.UUID(block_id),
            PaperBlock.paper_id == uuid.UUID(paper_id),
        )
    )
    block = result.scalar_one_or_none()
    if not block:
        raise NotFoundException("内容块")
    await db.delete(block)
    await db.commit()
    return {"detail": "内容块已删除"}


@router.post("/{paper_id}/blocks/reorder")
async def reorder_blocks(
    paper_id: str,
    request: ReorderBlocksRequest,
    user: User = Depends(require_permission(Permission.DOC_READ)),
    db: AsyncSession = Depends(get_db),
):
    await _get_paper_for_user(paper_id, user, db)
    for index, block_id in enumerate(request.block_ids):
        result = await db.execute(
            select(PaperBlock).where(
                PaperBlock.id == uuid.UUID(block_id),
                PaperBlock.paper_id == uuid.UUID(paper_id),
            )
        )
        block = result.scalar_one_or_none()
        if not block:
            raise NotFoundException(f"内容块 {block_id}")
        block.sort_order = index
    await db.commit()
    return {"detail": "排序已更新"}


@router.post("/{paper_id}/import", response_model=BlockResponse)
async def import_content(
    paper_id: str,
    request: ImportBlockRequest,
    user: User = Depends(require_permission(Permission.DOC_READ)),
    db: AsyncSession = Depends(get_db),
):
    paper = await _get_paper_for_user(paper_id, user, db)
    max_order_result = await db.execute(
        select(PaperBlock.sort_order)
        .where(PaperBlock.paper_id == paper.id)
        .order_by(PaperBlock.sort_order.desc())
        .limit(1)
    )
    max_order = max_order_result.scalar_one_or_none()
    next_order = (max_order + 1) if max_order is not None else 0

    block = PaperBlock(
        paper_id=paper.id,
        block_type=BlockType(request.block_type),
        content=request.content,
        sort_order=next_order,
        source_type=request.source_type,
        source_id=request.source_id,
    )
    db.add(block)
    await db.commit()
    await db.refresh(block)
    return block


@router.post("/{paper_id}/export/word")
async def export_word(
    paper_id: str,
    request: PaperExportRequest | None = None,
    user: User = Depends(require_permission(Permission.DOC_READ)),
    db: AsyncSession = Depends(get_db),
):
    paper = await _get_paper_for_user(paper_id, user, db)
    blocks_result = await db.execute(
        select(PaperBlock).where(PaperBlock.paper_id == paper.id).order_by(PaperBlock.sort_order)
    )
    blocks = list(blocks_result.scalars().all())
    content = _build_word_export(paper, blocks, request.watermark_text if request else None)
    filename = _download_filename(paper.title, "docx")
    return StreamingResponse(
        BytesIO(content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"},
    )


@router.post("/{paper_id}/export/pdf")
async def export_pdf(
    paper_id: str,
    request: PaperExportRequest | None = None,
    user: User = Depends(require_permission(Permission.DOC_READ)),
    db: AsyncSession = Depends(get_db),
):
    paper = await _get_paper_for_user(paper_id, user, db)
    blocks_result = await db.execute(
        select(PaperBlock).where(PaperBlock.paper_id == paper.id).order_by(PaperBlock.sort_order)
    )
    blocks = list(blocks_result.scalars().all())
    content = _build_pdf_export(paper, blocks, request.watermark_text if request else None)
    filename = _download_filename(paper.title, "pdf")
    return StreamingResponse(
        BytesIO(content),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"},
    )


@router.post("/{paper_id}/export/markdown")
async def export_markdown(
    paper_id: str,
    user: User = Depends(require_permission(Permission.DOC_READ)),
    db: AsyncSession = Depends(get_db),
):
    await _get_paper_for_user(paper_id, user, db)
    return {"detail": "Markdown 导出功能即将上线"}
